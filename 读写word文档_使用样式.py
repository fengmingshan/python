# -*- coding: utf-8 -*-
"""
Created on Thu Sep  5 17:16:28 2019

@author: Administrator
"""

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import os

data_path = 'd:/test/'

os.chdir(data_path)

doc = Document()

# 打印出所有的字体样式
styles = doc.styles
print("\n".join([s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]))


document = Document('d:/test/demo.docx')
for p in document.paragraphs:
    print(len(p.text))
    print(p.style.name)

# 用游程遍历段落
document = Document('d:/test/demo.docx')
for p in document.paragraphs:
    print('====')
    for r in p.runs:
        print(len(r.text))

# =============================================================================
# 修改字体
# =============================================================================
from docx.oxml.ns import qn

doc = Document()
paragraph doc.add_paragraph(u'一个人的命运啊，当然要靠自我奋斗')
# 打印所有styles
#for style in doc.styles:
#    print(style.name)
doc.styles['Normal'].font.name = u'华文楷体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
paragraph.style = doc.styles['Normal']

doc.save('4-0.docx')

# =============================================================================
# 不同字体大小
# =============================================================================
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


doc = Document()
for i in range(10):
    p = doc.add_paragraph(u'段落 %d' % i)
    style = doc.styles.add_style('UserStyle%d' %i,WD_STYLE_TYPE.PARAGRAPH)
    doc.styles['UserStyle%d'%i].font.name = u'华文楷体'
    doc.styles['UserStyle%d'%i]._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    style.font.size = Pt(i+20)
    p.style = style

doc.save('4-1.docx')


# =============================================================================
# 不同字体和不同字体颜色
# =============================================================================

from docx.shared import RGBColor
from docx.oxml.ns import qn

doc = Document()
p = doc.add_paragraph()
text_str = u'一个人的命运啊，当然要靠自我奋斗，但是也要考虑到历史的进程。'
for i,ch in enumerate(text_str):
    run = p.add_run(ch)
    font = run.font
    font.name = u'微软雅黑'
    # bug of python-docx
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    font.bold = (i % 2 == 0)
    font.italic = (i % 3 == 0)
    color = font.color
    color.rgb = RGBColor(i*10 % 200 + 55,i*20 % 200 + 55,i*30 % 200 + 55)

doc.save('4-2.docx')



