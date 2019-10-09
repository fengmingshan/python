# -*- coding: utf-8 -*-
"""
Created on Thu Sep  5 16:07:52 2019

@author: Administrator
"""


from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import os
from docx.shared import Pt

data_path = 'd:/_小程序/python读写word文档'

os.chdir(data_path)

document = Document()

# 修改文档字体为楷体
document.styles['Normal'].font.name = u'华文楷体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
document.styles['Normal'].font.size = Pt(20)
style_1 = document.styles['Normal']

# 添加一个新的字体
style_2 = document.styles.add_style('style_2', WD_STYLE_TYPE.PARAGRAPH)
document.styles['style_2'].font.name = u'微软雅黑'
document.styles['style_2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
style_2.font.size = Pt(20)

# 添加标题 等级0,并修改标题的字体，
# 注意修改标题字体必须用run（游标）来完成，不能直接修改。
title = document.add_heading('', level=0).add_run('4G话务周报')
title.font.name = u'华文楷体'
title._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')

# 添加第一段
head1 = document.add_heading('', level=1).add_run('一、基站数量')
head1.font.name = u'华文楷体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head1.font.size = Pt(20)

# 添加l两个小节
head2 = document.add_heading('', level=2).add_run('1.1 各县基站数量分析')
head2.font.name = u'华文楷体'
head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head2.font.size = Pt(16)


paragraph1 = document.add_paragraph('本周各县均无无新增基站', style=style_1)

head2 = document.add_heading('', level=2).add_run('1.2 各县超闲基站分析')
head2.font.name = u'华文楷体'
head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head2.font.size = Pt(16)

# 添加图片
document.add_picture('./pic/CQI优良比.png', width=Inches(2.25))

# 添加第二段
head1 = document.add_heading('', level=1).add_run('二、4G流量')
head1.font.name = u'华文楷体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head1.font.size = Pt(20)

paragraph2 = document.add_paragraph(
    '本周4G流量较上周大幅度增长，增长较多的有宣威、会泽、富源；下降的有麒麟、沾益、马龙。', style=style_2)

# 添加图片
document.add_picture('./pic/MR覆盖.png', width=Inches(2.25))

# 添加第三段
head1 = document.add_heading('', level=1).add_run('三、4G用户数')
head1.font.name = u'华文楷体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head1.font.size = Pt(20)


paragraph3 = document.add_paragraph(
    '本周4G用户数上周有所减少，增长的有宣威、会泽、富源；下降的有麒麟、沾益、马龙。', style=style_1)


# 添加图片
document.add_picture('./pic/RRC用户数.png', width=Inches(2.25))


document.add_page_break()

document.save('demo.docx')


#document = Document('./demo.docx')
#for p in document.paragraphs:
#    #print(len(p.text))
#    print(p.style.name)

