# -*- coding: utf-8 -*-
"""
Created on Mon Sep 30 16:23:08 2019
@author: Administrator
"""

from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches
from docx import Document
import pandas as pd
import os
import matplotlib.pyplot as plt
from math import ceil
plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号

# 导入读写word文档的库

path = 'D:/_小程序/感知工单分析报告'
pic_path = 'D:/_小程序/感知工单分析报告/pic/'
analyze_cell = '730448_6'

if not os.path.exists(pic_path):
    os.mkdir(pic_path)
os.chdir(path)

# 根据小区是L800还是L1800确定用户数超忙门限和流量超忙门限
L800_list = [17, 18, 19, 20, 21, 22,145, 146, 147, 148, 149,150]
L1800_list = [49, 50, 51, 52, 53, 54, 55, 56,177, 178, 179, 180, 181, 182]
L2100_list = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11,129, 130, 131, 132, 133, 134, 135, 136]

if int(analyze_cell.split('_')[1]) in L800_list:
    user_busy = 50
    throughput_busy = 1.5
else :
    user_busy = 200
    throughput_busy = 6

# =============================================================================
# 定义画图函数和一些判断函数
# =============================================================================
def draw_line_chart(df, time_col, data_col):
    x1 = df[time_col].values
    y1 = df[data_col].values
    plt.figure(figsize=(10, 5))
    plt.xticks(range(len(x1)), x1, fontsize=8)
    plt.plot(range(len(x1)), y1, label=data_col, linewidth=3,
             color='b', marker='o', markerfacecolor='yellow', markersize=6)
    for a, b in zip(range(len(x1)), y1):
        plt.text(a, b * 1.001, b, ha='center', va='bottom', fontsize=10)
    plt.xlabel('时间')
    plt.ylabel(data_col)
    plt.legend(loc='best')
    plt.title(data_col)
    #plt.legend(bbox_to_anchor=(1.01, 0.5), loc='upper left', borderaxespad=0.)
    plt.savefig(pic_path + data_col + '.png', format='png', dpi=400)
    plt.show('hold')


def draw_bar_chart(df, time_col, data_col):
    x2 = df[time_col].values
    y2 = df[data_col].values
    plt.figure(figsize=(15, 5))
    plt.bar(x2, y2, color='g', width=0.3, alpha=0.6, label=data_col)
    for x, y in zip(x2, y2):
        plt.text(x, y * 1.005, y, ha='center', va='bottom', fontsize=8)
    plt.xlabel('时间')
    plt.xticks(range(0, len(x2)), x2)
    plt.ylabel(data_col)
    # 将图例扩展到图片外部，x坐标1.05在图片外部0.05处，y坐标在图片中间，borderaxespad=0.表示扩展到图片外部
    #plt.legend(bbox_to_anchor=(1.01, 0.5), loc='upper left', borderaxespad=0.)
    plt.legend(loc='best')
    plt.title(data_col)
    plt.savefig(pic_path + data_col + ".png", format='png', dpi=200)
    plt.show('hold')

def mr_cqi_level(mr_good_rate):
    mr_good_rate = float(mr_good_rate.replace('%', ''))
    if mr_good_rate >= 95:
        return '非常好'
    elif mr_good_rate >= 90:
        return '良好'
    elif mr_good_rate >= 85:
        return '一般'
    elif mr_good_rate < 85:
        return '稍差'

def speed_level(user_speed):
    if user_speed >=  5:
        return '非常好'
    elif user_speed >= 3.2:
        return '良好'
    elif user_speed >= 2:
        return '正常'
    elif user_speed >= 1:
        return '稍差'
    elif user_speed < 1:
        return '很差'

def prb_level(prb_ratio):
    if prb_ratio >=  90:
        return '非常高，急需扩容'
    elif prb_ratio >= 70:
        return '很高，需要扩容'
    elif prb_ratio >= 50:
        return '达到超忙小区标准，需要继续观察'
    elif prb_ratio < 50:
        return '正常'

def coverage_problem_locate(mr_good_rate,cqi_rate):
    mr_good_rate = float(mr_good_rate.replace('%',''))
    cqi_rate = float(cqi_rate.replace('%',''))
    if mr_good_rate >=  90:
        coverage =  'MR覆盖良好，'
        if cqi_rate >=  90:
            cqi_coverage =  'CQI优良率正常。'
        else:
            cqi_coverage =  '但CQI优良率稍差，可能存在mod3干扰，需要检查邻区和PCI规划。'
    else:
        coverage =  'MR覆盖良好'
        if cqi_rate >=  90:
            cqi_coverage =  'CQI优良率正常，无需优化。'
        else:
            cqi_coverage =  'CQI优良率稍差，需要通过优化增强覆盖。'
    coverage_problem = coverage + cqi_coverage
    return coverage_problem

def load_problem_locate(prb_ratio,max_uesrs,cell_throughput):
    if prb_ratio <  50:
        cell_load = '小区负荷正常，未达到超忙标准。'
    elif prb_ratio >  50 and max_uesrs > user_busy:
        cell_load = '小区达到多用户超忙门限，需要进行扩容建设才能解决。'
    elif prb_ratio >  50 and cell_throughput > throughput_busy:
        cell_load = '小区达到大流量超忙门限，需要进行扩容建设才能解决.'
    return cell_load


all_files = os.listdir(path)
KPI_file = [x for x in all_files if 'KPI' in x][0]
MR_file = [x for x in all_files if 'MR' in x][0]

# =============================================================================
# 话务分析
# =============================================================================
df_kpi = pd.read_csv(KPI_file, skiprows=5, engine='python')

df_kpi['hour'] = df_kpi['开始时间'].map(lambda x: x.split(' ')[1][:2])
hour_list = list(set(df_kpi['hour']))
# 计算实际忙时
for hour in hour_list:
    df_tmp = df_kpi[df_kpi['hour'] == hour]
    ind = df_kpi.groupby(by='hour', as_index=True)[
        '最大RRC连接用户数_1'].agg(sum).idxmax()
df_busy = df_kpi[df_kpi['hour'] == ind]
df_busy['下行PRB平均占用率_1'] = df_busy['下行PRB平均占用率_1'].map(
    lambda x: float(x.replace('%', '')))
df_busy['E-RAB掉线率_1'] = df_busy['E-RAB掉线率_1'].map(
    lambda x: float(x.replace('%', '')))
df_busy['下行平均激活用户数_1'] = df_busy['下行平均激活用户数_1'].map(
    lambda x: ceil(x))
df_busy['CQI>=7占比'] = df_busy['CQI>=7占比'].map(
    lambda x: float(x.replace('%', '')))
df_busy['空口下行用户面流量（MByte）_1477070755617-11'] = df_busy['空口下行用户面流量（MByte）_1477070755617-11'].map(
    lambda x: float(x.replace(',', '')))

df_busy['time'] = df_busy['开始时间'].map(lambda x: x[5:])
df_busy.rename(columns={'CQI>=7占比': 'CQI优良比'}, inplace=True)

cqi_rate = str(round(df_busy['CQI优良比'].mean(), 2)) + '%'
cqi_level = mr_cqi_level(cqi_rate)

prb_ratio = round(df_busy['CQI优良比'].mean(),2)
prb_level = prb_level(prb_ratio)

user_speed = round(df_busy['分QCI用户体验下行平均速率（Mbps）_1'].mean(),2)
user_speed_level = speed_level(user_speed)

max_uesrs = df_busy['最大RRC连接用户数_1'].max()
avg_uesrs = ceil(df_busy['平均RRC连接用户数_1'].mean())
max_act_uesrs = df_busy['最大激活用户数_1'].max()
avg_act_uesrs = ceil(df_busy['平均激活用户数_1'].mean())

cell_throughput = round(df_busy['空口下行用户面流量（MByte）_1477070755617-11'].mean()/1024,2)

draw_line_chart(df_busy, 'time', '下行PRB平均占用率_1')
draw_line_chart(df_busy, 'time', 'E-RAB掉线率_1')
draw_line_chart(df_busy, 'time', 'CQI优良比')
draw_bar_chart(df_busy, 'time', '最大RRC连接用户数_1')
draw_bar_chart(df_busy, 'time', '平均RRC连接用户数_1')
draw_bar_chart(df_busy, 'time', '最大激活用户数_1')
draw_bar_chart(df_busy, 'time', '平均激活用户数_1')
draw_bar_chart(df_busy, 'time', '分QCI用户体验下行平均速率（Mbps）_1')

# 画用户数图
df_busy[['最大RRC连接用户数_1','平均RRC连接用户数_1','最大激活用户数_1','平均激活用户数_1']].plot(figsize=(15, 5),legend=True)
#plt.legend(bbox_to_anchor=(1.01, 0.5), loc='upper left', borderaxespad=0.)
plt.savefig('./pic/用户数.png', format='png', dpi=200)
plt.show()
# =============================================================================
# MR分析
# =============================================================================
df_mr = pd.read_csv(MR_file, engine='python')
df_mr['cell_ind'] = df_mr['NAME'].apply(
    lambda x: x.split('_')[0] + '_' + x.split('_')[1])
df_mr = df_mr[df_mr['cell_ind'] == analyze_cell]
df_mr['points'] = df_mr['|≥-105dBm采样点'] + df_mr['|≥-110dBm采样点'] + \
    df_mr['|≥-115dBm采样点'] + df_mr['|≥-120dBm采样点'] + df_mr['|≥负无穷采样点']
df_mr['good_points'] = df_mr['|≥-105dBm采样点'] + df_mr['|≥-110dBm采样点']
df_mr['MR覆盖率'] = round(df_mr['good_points'] / df_mr['points'] * 100, 2)
average_rsrp = df_mr.loc[:, '平均RSRP（dBm）'].values[0]

mr_good_rate = str(df_mr.loc[:, 'MR覆盖率'].values[0]) + '%'
mr_level = mr_cqi_level(mr_good_rate)
coverage_problem = coverage_problem_locate(mr_good_rate,cqi_rate)
cell_load = load_problem_locate(prb_ratio,max_uesrs,cell_throughput)

# 对MR的字段进行重命名
df_mr.set_index('NAME', drop=False, inplace=True)
df_mr.rename(
    columns={
        '|≥-105dBm采样点': '≥-105dBm',
        '|≥-110dBm采样点': '≥-110dBm',
        '|≥-115dBm采样点': '≥-115dBm',
        '|≥-120dBm采样点': '≥-120dBm',
        '|≥负无穷采样点': '≥负无穷'},
    inplace=True)

# 画MR覆盖图
df_mr[['≥-105dBm', '≥-110dBm', '≥-115dBm', '≥-120dBm', '≥负无穷']
      ].T.plot(kind='bar', figsize=(10, 5), legend=True,title='MR覆盖分布')
plt.savefig('./pic/MR覆盖情况' + ".png", format='png', dpi=200)
plt.show('hold')

analyze_cell_name = analyze_cell + '_' + \
    df_mr.loc[:, 'NAME'].values[0].split('QJ')[1].split('_')[0]

# =============================================================================
# 写入到word文档
# =============================================================================
document = Document()
# 定义文档默认字体style_1为楷体
document.styles['Normal'].font.name = u'华文楷体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
document.styles['Normal'].font.size = Pt(14)
style_1 = document.styles['Normal']

# 添加文档标题 等级0
title = document.add_heading('', level=0).add_run(analyze_cell_name + '感知质差分析')
title.font.name = u'华文楷体'
title._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
title.font.size = Pt(25)

# 添加第一段标题
head1 = document.add_heading('', level=1).add_run('一、小区覆盖分析')
head1.font.name = u'华文楷体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head1.font.size = Pt(20)

# 添加第一段的1小节标题
head2 = document.add_heading('', level=2).add_run('1.1 MR覆盖情况')
head2.font.name = u'华文楷体'
head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head2.font.size = Pt(16)

# 添加MR覆盖率图片
document.add_picture('./pic/MR覆盖情况' + ".png", width=Inches(6))
# 添加MR覆盖率描述
document.add_paragraph(
    analyze_cell_name +
    '小区: ''MR覆盖率:{mr_rate},MR覆盖指标{mr_level}。'.format(
        mr_rate=mr_good_rate,
        mr_level = mr_level),
    style=style_1)

# 添加第一段的2小节标题
head2 = document.add_heading('', level=2).add_run('1.2 CQI优良比')
head2.font.name = u'华文楷体'
head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head2.font.size = Pt(16)

# 添加CQI图片
document.add_picture('./pic/CQI优良比' + ".png", width=Inches(6))
# 添加CQI文字描述
document.add_paragraph(
    '小区: ' + analyze_cell_name +
    '平均CQI覆盖率:{cqi_rate},CQI覆盖率指标{cqi_level}。'.format(
        cqi_rate=cqi_rate,
        cqi_level=cqi_level),
    style=style_1)

# 添加第一段的3小节标题
head2 = document.add_heading('', level=2).add_run('1.3 用户体验速率')
head2.font.name = u'华文楷体'
head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head2.font.size = Pt(16)

# 添加速率图片
document.add_picture('./pic/分QCI用户体验下行平均速率（Mbps）_1' + ".png", width=Inches(6))
# 添加速率文字描述
document.add_paragraph(
    '小区: '+ analyze_cell_name +
    '平均用户体验速率:{user_speed},用户体验速率{user_speed_level}。'.format(
        user_speed = user_speed,
        user_speed_level = user_speed_level),
    style=style_1)



# 添加第二段标题
head1 = document.add_heading('', level=1).add_run('二、小区话务负荷分析')
head1.font.name = u'华文楷体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head1.font.size = Pt(20)

# 添加第一段的1小节标题
head2 = document.add_heading('', level=2).add_run('2.1 用户数分析')
head2.font.name = u'华文楷体'
head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head2.font.size = Pt(16)

# 添加用户数图片
document.add_picture('./pic/用户数' + ".png", width=Inches(6))
# 添加用户数说明
document.add_paragraph(
    '小区: ' +
    analyze_cell_name +
    '忙时最大RRC连接用户数{max_uesrs}个，平均RRC连接用户数{avg_uesrs}个,最大激活用户数{max_act_uesrs}个,平均激活用户数{avg_act_uesrs}个。'.format(
        max_uesrs = max_uesrs,
        avg_uesrs = avg_uesrs,
        max_act_uesrs = max_act_uesrs,
        avg_act_uesrs = avg_act_uesrs),
    style=style_1)

# 添加第一段的2小节标题
head2 = document.add_heading('', level=2).add_run('2.2 PRB利用率分析')
head2.font.name = u'华文楷体'
head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head2.font.size = Pt(16)

# 添加用户数图片
document.add_picture('./pic/下行PRB平均占用率_1' + ".png", width=Inches(6))
# 添加用户数说明
document.add_paragraph(
    '小区: ' +
    analyze_cell_name +
    'PRB下行利用率{prb_ratio}%，属于{prb_level}。'.format(
        prb_ratio = prb_ratio,
        prb_level = prb_level),
    style=style_1)

# 添加第三段标题
head1 = document.add_heading('', level=1).add_run('三、问题定位')
head1.font.name = u'华文楷体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
head1.font.size = Pt(20)
# 添加文字说明
document.add_paragraph(
    '小区: ' +
    analyze_cell_name +
    ',覆盖情况：{coverage_problem}。小区负荷情况：{cell_load}。'.format(
        coverage_problem = coverage_problem,
        cell_load = cell_load),
    style=style_1)

# 保存文档
document.add_page_break()
document.save(analyze_cell_name + '感知质差分析报告.docx')
